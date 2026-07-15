#!/usr/bin/env python3
"""Convert a markdown report to a professionally styled DOCX.

Usage:
    gen_docx.py <input.md> [output.docx]

If output is omitted, writes to <input_stem>.docx in the same directory.
Requires: pip install python-docx
"""

import os
import re
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ─── Color palette ───
HEADER_BG = "2F5496"
HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
HEADING_COLOR = RGBColor(0x2F, 0x54, 0x96)
ALT_ROW_BG = "F2F2F2"
HIGHLIGHT_BG = "FFF2CC"
BORDER_COLOR = "BFBFBF"
CODE_COLOR = RGBColor(0xC0, 0x39, 0x2B)
BODY_FONT = "Calibri"
CODE_FONT = "Consolas"
BODY_SIZE = Pt(10)
TABLE_FONT_SIZE = Pt(9)


def set_cell_margins(table, top=40, bottom=40, start=60, end=60):
    """Set uniform cell margins on the entire table. Values in twips."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    margins = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{start}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{end}" w:type="dxa"/>'
        f'</w:tblCellMar>'
    )
    tblPr.append(margins)


def set_table_borders(table):
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'</w:tblBorders>'
    )
    table._tbl.tblPr.append(borders)


def shade_cells(cells, color):
    for cell in cells:
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)


def set_row_height(row, height_pt):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{int(height_pt * 20)}" w:hRule="atLeast"/>'
    )
    trPr.append(trHeight)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    pPr = paragraph._p.get_or_add_pPr()
    spacing_attrs = f'w:before="{before}" w:after="{after}"'
    if line:
        spacing_attrs += f' w:line="{line}" w:lineRule="auto"'
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} {spacing_attrs}/>')
    existing = pPr.find(f'{{{nsdecls("w").split(chr(34))[1]}}}spacing')
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)


def add_runs_with_formatting(paragraph, text):
    """Parse **bold** and `code` inline and add styled runs."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = BODY_FONT
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt(9)
            run.font.color.rgb = CODE_COLOR
        else:
            run = paragraph.add_run(part)
            run.font.name = BODY_FONT


def add_image(doc, path, width=6.2):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        p = doc.paragraphs[-1]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, before=60, after=120)
    else:
        doc.add_paragraph(f"[Image not found: {path}]")


def parse_table_block(lines):
    headers = [c.strip() for c in lines[0].strip('|').split('|')]
    rows = []
    for line in lines[2:]:
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    return headers, rows


def add_styled_table(doc, headers, rows):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    set_table_borders(table)
    set_cell_margins(table)

    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        add_runs_with_formatting(p, h)
        for run in p.runs:
            run.font.size = TABLE_FONT_SIZE
            run.font.color.rgb = HEADER_FG
            run.font.bold = True
            run.font.name = BODY_FONT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_paragraph_spacing(p, before=20, after=20)

    shade_cells(table.rows[0].cells, HEADER_BG)
    set_row_height(table.rows[0], 14)

    for r_idx, row_data in enumerate(rows):
        has_bold = any('**' in c for c in row_data)
        row = table.rows[r_idx + 1]

        for c_idx, val in enumerate(row_data):
            if c_idx < n_cols:
                cell = row.cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                add_runs_with_formatting(p, val)
                for run in p.runs:
                    run.font.size = TABLE_FONT_SIZE
                    run.font.name = BODY_FONT
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_paragraph_spacing(p, before=15, after=15)

        if has_bold:
            shade_cells(row.cells, HIGHLIGHT_BG)
        elif r_idx % 2 == 1:
            shade_cells(row.cells, ALT_ROW_BG)

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=60)

    return table


def setup_styles(doc):
    """Configure document-wide styles."""
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    configs = [
        (0, "Title", Pt(22), True),
        (1, "Heading 1", Pt(16), True),
        (2, "Heading 2", Pt(13), True),
        (3, "Heading 3", Pt(11), True),
    ]
    for _, style_name, size, bold in configs:
        if style_name in doc.styles:
            s = doc.styles[style_name]
            s.font.name = BODY_FONT
            s.font.size = size
            s.font.bold = bold
            s.font.color.rgb = HEADING_COLOR
            pf = s.paragraph_format
            pf.space_before = Pt(14)
            pf.space_after = Pt(6)

    if "List Bullet" in doc.styles:
        lb = doc.styles["List Bullet"]
        lb.font.name = BODY_FONT
        lb.font.size = BODY_SIZE
        lb.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        pf = lb.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.2


def convert_md_to_docx(md_path, out_path):
    md_dir = os.path.dirname(os.path.abspath(md_path))

    with open(md_path, 'r') as f:
        lines = f.readlines()

    doc = Document()
    setup_styles(doc)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        if re.match(r'^-{3,}\s*$', line):
            i += 1
            continue

        m = re.match(r'^(#{1,4})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            heading_level = 0 if level == 1 else level - 1
            h = doc.add_heading(level=min(heading_level, 3))
            add_runs_with_formatting(h, text)
            for run in h.runs:
                run.font.color.rgb = HEADING_COLOR
            i += 1
            continue

        m = re.match(r'^!\[.*?\]\((.+?)\)\s*$', line)
        if m:
            img_path = os.path.join(md_dir, m.group(1))
            add_image(doc, img_path)
            i += 1
            continue

        if '|' in line and i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1].rstrip('\n')):
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].rstrip('\n'))
                i += 1
            headers, rows = parse_table_block(table_lines)
            add_styled_table(doc, headers, rows)
            continue

        m = re.match(r'^- (.+)$', line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_formatting(p, m.group(1))
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        para_parts = []
        while i < len(lines):
            l = lines[i].rstrip('\n')
            if not l.strip():
                break
            if re.match(r'^#{1,4}\s', l):
                break
            if re.match(r'^!\[', l):
                break
            if re.match(r'^-{3,}\s*$', l):
                break
            if '|' in l and i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1].rstrip('\n')):
                break
            if re.match(r'^- ', l):
                break
            para_parts.append(l)
            i += 1
        text = ' '.join(para_parts)
        p = doc.add_paragraph()
        add_runs_with_formatting(p, text)
        continue

    doc.save(out_path)
    print(f"Report saved to: {out_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: gen_docx.py <input.md> [output.docx]")
        sys.exit(1)
    md_file = sys.argv[1]
    if len(sys.argv) > 2:
        out_file = sys.argv[2]
    else:
        out_file = os.path.splitext(md_file)[0] + ".docx"
    convert_md_to_docx(md_file, out_file)
